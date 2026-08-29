#!/usr/bin/env python3
"""Apply the corrections of worklog/05-reference-verification.md to the four course documents.

Run from the repo root. Idempotent: every replacement is keyed on the *old* text, so a second run
finds nothing to change. Originals are kept in worklog/course-originals/.
"""
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[2]
COURSE = ROOT / "course"
DATE = "29 August 2026"
GUIDELINE_URL = "https://www.buildwise.be/media/srahyb00/na-ec-7-beschoeiingen-2022-final-nl.pdf"
ANCHORS_URL = "https://www.buildwise.be/media/mxyludtd/nb-ec-7-grondankers.pdf"
MANUALS_URL = "https://files.seequent.com/PLAXIS/Manuals/PLAXIS_2D/English/"
EC7_2024_URL = "https://www.buildwise.be/en/publications/standards-regulations/en-1997-1-2024-en/"
EC3_5_2025_URL = "https://www.buildwise.be/en/publications/standards-regulations/en-1993-5-2025-en/"

GUIDELINE_EN = (
    "Normalisatiecommissie NBN E25007 “Eurocode 7” (WTCB/Buildwise – SECO) (2022). "
    "Richtlijnen voor de toepassing van de Eurocode 7 in België volgens NBN EN 1997-1 ANB – "
    "Het grondmechanische ontwerp van ingebedde kerende constructies: beschoeiingen. March 2022, 39 pp. "
    "Brussels: Buildwise / BGGG-GBMS. " + GUIDELINE_URL
)

changes = []


def set_paragraph_text(p, text):
    """Replace the whole paragraph text, keeping the formatting of the first run."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def replace_in_paragraph(p, old, new, label):
    """Substring replacement that survives run boundaries (formatting of the first affected run is kept)."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    # try single-run replacement first (keeps every run's formatting)
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            changes.append(label)
            return True
    # spans several runs: rebuild from the first affected run onwards
    start = full.index(old)
    pos = 0
    first = None
    for i, r in enumerate(p.runs):
        if pos + len(r.text) > start:
            first = i
            break
        pos += len(r.text)
    head = "".join(r.text for r in p.runs[:first])
    tail = full[len(head):].replace(old, new, 1)
    p.runs[first].text = tail
    for r in p.runs[first + 1:]:
        r.text = ""
    changes.append(label)
    return True


def replace_ref(p, number, new_body, label):
    """Numbered reference '[n] body' held in one or two runs ('[n] ' + body)."""
    full = "".join(r.text for r in p.runs)
    prefix = f"[{number}]"
    if not full.startswith(prefix):
        print(f"  ! expected {prefix} at paragraph, found: {full[:60]!r}")
        return False
    if len(p.runs) >= 2 and p.runs[0].text.strip() == prefix:
        p.runs[1].text = new_body + (" " if full.endswith(" ") else "")
        for r in p.runs[2:]:
            r.text = ""
    else:
        set_paragraph_text(p, f"{prefix} {new_body}")
    changes.append(label)
    return True


def find(doc, needle, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and needle in p.text:
            return i
    raise KeyError(needle)


# --------------------------------------------------------------------------------------------
def sheet_pile_manual():
    f = COURSE / "Sheet_Pile_Retaining_Walls_Manual_EC7_PLAXIS_v24.docx"
    d = docx.Document(f)
    ps = d.paragraphs
    r0 = find(d, "References", start=200)
    ref = {int(ps[i].text.strip()[1:ps[i].text.index("]")]): ps[i]
           for i in range(r0 + 1, r0 + 18) if ps[i].text.strip().startswith("[")}
    ws = " [working source]"
    replace_in_paragraph(ref[1], "pending the Belgian ANB.", "pending the Belgian ANB. " + EC7_2024_URL + ".", "SP[1] url")
    replace_ref(ref[3], 3, GUIDELINE_EN + ws, "SP[3] issuer/title/url")
    replace_in_paragraph(ref[5], "pending the new Belgian ANB.", "pending the new Belgian ANB. " + EC3_5_2025_URL + ".", "SP[5] url")
    replace_ref(ref[6], 6, "NBN EN 12063:2024. Execution of special geotechnical work — Sheet pile walls, combined pile walls, "
                "high modulus walls. Brussels: NBN (EN 12063:2024, CEN, May 2024; supersedes EN 12063:1999)." + ws, "SP[6] title")
    replace_ref(ref[7], 7, "Seequent / Bentley Systems. Material datasets for plates: sheet pile wall in bending. PLAXIS Knowledge Base "
                "article KB0110039 (formerly Bentley Communities wiki 45929); worked sheet-pile Plate data set." + ws, "SP[7] title/KB")
    replace_ref(ref[8], 8, "Seequent, The Bentley Subsurface Company (2025). PLAXIS 2D 2025.1 Reference Manual, chapter “Interfaces” "
                "(interface strength, stiffness and permeability; last updated 24 September 2025). " + MANUALS_URL + ws, "SP[8] manual")
    replace_ref(ref[9], 9, "Seequent / Bentley Systems. End bearing of plates. PLAXIS Knowledge Base article KB0110231 "
                "(formerly Bentley Communities wiki 46016, 2016)." + ws, "SP[9] KB")
    replace_ref(ref[10], 10, "Seequent, The Bentley Subsurface Company (2025). PLAXIS 2D 2025.1 Reference Manual, “Interfaces tabsheet — "
                "Groundwater”: cross-permeability options (impermeable, semi-permeable, fully permeable). " + MANUALS_URL + ws, "SP[10] manual")
    replace_ref(ref[11], 11, "Seequent, The Bentley Subsurface Company (2025). PLAXIS 2D 2025.1 Reference Manual, “Interfaces tabsheet”: "
                "definition of hydraulic resistance d/k and drainage conductivity d·k. " + MANUALS_URL + ws, "SP[11] manual")
    replace_ref(ref[12], 12, "Seequent, The Bentley Subsurface Company (2025). PLAXIS 2D 2025.1 Tutorial Manual, Tutorial 6: Dry excavation using a "
                "tie back wall (last updated 24 September 2025). Node-to-node anchor for the free length, Embedded Beam Row for the grout body "
                "(T_skin = 400 kN/m). " + MANUALS_URL + "PLAXIS_2D_1_Tutorial%20Manual.pdf" + ws, "SP[12] tutorial")
    replace_ref(ref[14], 14, "Simpson, B. and Powrie, W. (2001). Embedded retaining walls: theory, practice and understanding. Proceedings of the "
                "15th International Conference on Soil Mechanics and Geotechnical Engineering, Istanbul, Vol. 4, Balkema, pp. 2505–2524." + ws, "SP[14] pages")
    replace_ref(ref[15], 15, "ArcelorMittal Sheet Piling (2016, rev. 2022). Piling Handbook, 9th edition (2016; chapters 1 and 3 revised 2022), "
                "aligned with EN 1997-1 and EN 1993-5. Luxembourg: ArcelorMittal Commercial RPS." + ws, "SP[15] edition")
    replace_in_paragraph(ref[16], "official current manual index.", "official current manual index (2025.1 manuals). " + MANUALS_URL, "SP[16] url")
    replace_ref(ref[17], 17, "Buildwise werkgroep “Beschoeiingen” / Normalisatiecommissie NBN E25007 (2024). Richtlijnen voor de toepassing van de "
                "Eurocode 7 in België volgens NBN EN 1997-1 ANB – Deel 3: Het grondmechanische ontwerp van voorgespannen groutankers. "
                "March 2024. " + ANCHORS_URL + ws, "SP[17] title")
    # body: the unverifiable verbatim attribution about Prevent punching
    i = find(d, "Bentley explicitly states that Prevent punching")
    replace_in_paragraph(ps[i],
        "Bentley explicitly states that Prevent punching shall not be used for sheet-pile walls because it creates an artificial elastic zone and does not represent physical end bearing [9].",
        "The Prevent punching option only generates a small elastic zone around the plate tip [9]; it is a numerical device, not a physical end bearing, "
        "and is therefore set to No for sheet-pile walls in this manual.", "SP body prevent-punching")
    i = find(d, "Bentley’s official sheet-pile example derives")
    replace_in_paragraph(ps[i], "Bentley’s official sheet-pile example", "Bentley’s Knowledge Base sheet-pile example", "SP body [7] wording")
    # revision line
    replace_in_paragraph(ps[8], "Revision 1 — 27 August 2026",
                         "Revision 1b — " + DATE + " (cross-reference to the soldier-pile chapter in §1; reference list verified online and corrected)", "SP revision")
    d.save(f)


def brinch_hansen_chapter():
    f = COURSE / "Brinch_Hansen_Tlat_Soldier_Pile_Walls_Textbook_Section.docx"
    d = docx.Document(f)
    ps = d.paragraphs
    r0 = find(d, "References", start=200)
    ref = {int(ps[i].text.strip()[1:ps[i].text.index("]")]): ps[i]
           for i in range(r0 + 1, r0 + 14) if ps[i].text.strip().startswith("[")}
    replace_ref(ref[4], 4, GUIDELINE_EN, "BH[4] issuer/title/url")
    replace_ref(ref[7], 7, "Seequent, The Bentley Subsurface Company (2025). PLAXIS 2D 2024.3 Reference Manual (2024.3 released 2025; current manuals: "
                "PLAXIS 2D 2025.1, 24 September 2025). " + MANUALS_URL, "BH[7] year/publisher")
    replace_ref(ref[8], 8, "Seequent / Bentley Systems (2025). 2D Analysis of an Anchored Soldier Pile Wall. PLAXIS Knowledge Base tutorial KB0045693 (2025).",
                "BH[8] KB")
    replace_ref(ref[9], 9, "Seequent / Bentley Systems (2026). PLAXIS 2D 2025.1 Release Notes, section “New in PLAXIS 2D 2025.1.2” (February 2026), "
                "fixed issue [1749896]: “The strength of Tlat for embedded beams is now correctly reduced in a safety calculation when Apply Strength "
                "Reduction marked active on this structural element.” Knowledge Base article KB0047805.", "BH[9] year/KB")
    replace_ref(ref[10], 10, "Buildwise (2024–2026). EN 1997-1:2024 – Eurocode 7: Geotechnical design – Part 1: General rules — status page: "
                "available for information only, not yet a Belgian standard pending the Belgian National Annex; NBN EN 1997-1:2005 + A1:2014 remain "
                "applicable. Status to be rechecked for every project. " + EC7_2024_URL, "BH[10] status page")
    # revision line + source note
    i = find(d, "Revision 2 — 29 August 2026")
    replace_in_paragraph(ps[i], "Revision 2 — 29 August 2026 (Sections 4.5–4.6 added)",
                         "Revision 2 — 29 August 2026 (Sections 4.5–4.6 added; reference list verified online and corrected)", "BH revision")
    d.save(f)


def vibratory_chapter():
    f = COURSE / "Vibratory_Pile_Installation_Manual_Eurocode_Course_Chapter.docx"
    d = docx.Document(f)
    ps = d.paragraphs
    i = find(d, "Van Rompaey, D., Legrand, C. and Holeyman, A.", start=900)
    replace_in_paragraph(ps[i], ", 15, 533-542. https://doi.org/10.2495/SD950601.",
        ", 14 (Proceedings of the 7th International Conference on Soil Dynamics and Earthquake Engineering, Chania, Crete, 24-26 May 1995), 533-542. "
        "Southampton: WIT Press. ISSN 1743-3509. https://www.witpress.com/Secure/elibrary/papers/SD95/SD95060FU.pdf", "VB Van Rompaey vol/DOI")
    i = find(d, "Holeyman, A. E. (2002)", start=900)
    replace_in_paragraph(ps[i], "Soil behavior under vibratory driving.", "Soil behaviour under vibratory driving.", "VB Holeyman 2002 title")
    replace_in_paragraph(ps[i], "Keynote lecture in ", "Keynote lecture in Holeyman, A., Vanden Berghe, J.-F. and Charue, N. (eds), ", "VB Holeyman 2002 editors")
    replace_in_paragraph(ps[i], ", 3-20. Lisse:", ", Louvain-la-Neuve, 9-10 September 2002, 3-20. Lisse:", "VB Holeyman 2002 venue")
    i = find(d, "Holeyman, A. and Whenham, V.", start=900)
    replace_in_paragraph(ps[i], "Critical Review of the Hypervib1 Model to Assess Vibratory Pile Drivability.",
                         "Critical Review of the Hypervib1 Model to Assess Pile Vibro-Drivability.", "VB Whenham title")
    replace_in_paragraph(ps[i], ", 35, 1933-1951.", ", 35(5), 1933-1951.", "VB Whenham issue")
    i = find(d, "European Commission, Joint Research Centre.", start=900)
    replace_in_paragraph(ps[i], "European Commission, Joint Research Centre.",
        "European Commission, Joint Research Centre – Bogusz, W., Caplane, C., Hard, D., Idda, K., Ingram, P., Kanty, P., Kushwaha, A., Nayrand, N., "
        "Sand, O., Sciarretta, F., Tsitsas, G. and Vogt, H.", "VB JRC authors")
    replace_in_paragraph(ps[i], "Implementation of Design during Execution and Service Life",
        "Implementation of Design during Execution & Service Life – Guidelines for the application of the 2nd generation of Eurocode 7: Geotechnical design",
        "VB JRC title")
    replace_in_paragraph(ps[i], ". Publications Office of the European Union. https://doi.org/10.2760/9211877.",
        ". JRC Technical Report JRC139606, EUR 40128. Luxembourg: Publications Office of the European Union. ISBN 978-92-68-22436-6. https://doi.org/10.2760/8383117.",
        "VB JRC DOI")
    i = find(d, "Nicholson, D., Tse, C.-M. and Penny, C.", start=900)
    replace_in_paragraph(ps[i], "(2000).", "(1999).", "VB CIRIA year")
    replace_in_paragraph(ps[i], "CIRIA Report R185. London: CIRIA.", "CIRIA Report R185. London: CIRIA. 214 pp.", "VB CIRIA pages")
    # in-text mentions of the CIRIA year, if any
    for p in ps[:900]:
        replace_in_paragraph(p, "Nicholson et al. (2000)", "Nicholson et al. (1999)", "VB in-text CIRIA year")
        replace_in_paragraph(p, "Nicholson, Tse and Penny (2000)", "Nicholson, Tse and Penny (1999)", "VB in-text CIRIA year")
    i = find(d, "Course edition 1.0 - standards status checked")
    replace_in_paragraph(ps[i], "Course edition 1.0 - standards status checked 29 August 2026",
                         "Course edition 1.0a - standards status and reference list checked online " + DATE, "VB revision")
    i = find(d, "Standards are copyrighted documents.", start=900)
    replace_in_paragraph(ps[i], "Standards are copyrighted documents.",
        "Reference list verified online on " + DATE + " (DOI resolution, publisher and catalogue records); the volume of Van Rompaey et al. (1995), "
        "the title of Holeyman and Whenham (2017), the JRC report identifiers and the CIRIA R185 year were corrected in this edition. "
        "Standards are copyrighted documents.", "VB reference-use note")
    d.save(f)


def rekennota():
    f = COURSE / "Rekennota_beschoeiing_berlinerwand_HEA180.docx"
    d = docx.Document(f)
    ps = d.paragraphs
    R = [
        ("NBN EN 1990 ANB:2013 - Eurocode 0: Grondslagen van het constructief ontwerp. Nationale Bijlage",
         "NBN EN 1990 ANB:2021 - Eurocode 0: Grondslagen van het constructief ontwerp. Nationale Bijlage (vervangt ANB:2013)"),
        ("NBN EN 1991-1-1:2002 - Eurocode 1: Belastingen op constructies. Deel 1-1: Algemene belastingen",
         "NBN EN 1991-1-1:2002 - Eurocode 1: Belastingen op constructies. Deel 1-1: Algemene belastingen (incl. AC:2009)"),
        ("NBN EN 1993-1-1 ANB:2010 - Belgische Nationale Bijlage bij EN 1993-1-1",
         "NBN EN 1993-1-1 ANB:2018 - Belgische Nationale Bijlage bij EN 1993-1-1 (vervangt ANB:2010)"),
        ("NBN EN ISO 22476-1:2013 - Geotechnisch onderzoek en beproeving. Veldproeven. Deel 1: Elektrische sondering met en zonder waterspanningsmeting",
         "NBN EN ISO 22476-1:2023 - Geotechnisch onderzoek en beproeving. Veldproeven. Deel 1: Elektrische sondering met en zonder waterspanningsmeting "
         "(vervangt NBN EN ISO 22476-1:2012 + AC:2013; sonderingen uitgevoerd vóór 2023 vallen onder de editie 2012)"),
        ("NBN EN 12063:1999 - Uitvoering van bijzonder geotechnisch werk. Damwanden",
         "NBN EN 12063:2024 - Uitvoering van bijzonder geotechnisch werk. Damwanden, combiwanden en wanden met hoge stijfheid (vervangt NBN EN 12063:1999). "
         "Berlinerwanden vallen buiten het toepassingsgebied van deze norm; zij wordt hier enkel als referentie voor de uitvoeringstoleranties van "
         "de stalen profielen aangehaald"),
        ("beschoeiingen. Maart 2022. Hierna aangeduid",
         "beschoeiingen. Maart 2022, 39 p. " + GUIDELINE_URL + ". Hierna aangeduid"),
        ("Standaardbestek 260 voor kunstwerken en waterbouw (SB260), Vlaamse Overheid, MOW. Hoofdstuk 21 - Geotechniek.",
         "Standaardbestek 260 voor kunstwerken en waterbouw (SB260), Vlaamse Overheid, MOW, «versie in te vullen» (versie 2.0, maart 2018, incl. "
         "errata en aanvullingen 2022; nieuwe editie verplicht vanaf 1 januari 2026). Hoofdstuk 21 - Geotechniek."),
        ("Standaardbestek 250 voor de wegenbouw in het Vlaamse Gewest, Agentschap Wegen en Verkeer.",
         "Standaardbestek 250 voor de wegenbouw in het Vlaamse Gewest, Agentschap Wegen en Verkeer, «versie in te vullen» (nieuwe editie verplicht "
         "vanaf 1 januari 2026)."),
        ("WTCB/Buildwise Infofiche 56.2 - Berlijnse wanden. Type 2: beschottingen aanbrengen tijdens de uitgraving (2012).",
         "WTCB/Buildwise Infofiche 56.2 - Berlijnse wanden. Type 2: beschottingen aanbrengen vóór de uitgraving (juli 2012)."),
        ("WTCB/Buildwise Infofiche 56.1 - Berlijnse wanden. Type 1: beschottingen aanbrengen tijdens de uitgraving (2012).",
         "WTCB/Buildwise Infofiche 56.1 - Berlijnse wanden. Type 1: beschottingen aanbrengen tijdens de uitgraving (juli 2012)."),
        ("BGGG (2012). Standaardprocedures voor geotechnisch onderzoek. Sonderingen, Deel 1: Planning, uitvoering en rapportering.",
         "BGGG-GBMS (2016). Standaardprocedures voor geotechnisch onderzoek: Sonderingen. Deel 1: Planning, uitvoering en rapportering "
         "(BGGG-CPT-pt1-2016, 14 juli 2016; Deel 2: Interpretatie, 27 april 2017). www.bggg-gbms.be"),
        ("CUR-rapport 2003-7. Bepaling geotechnische parameters. Stichting CUR, Gouda.",
         "CUR-rapport 2003-7 (2003). Bepaling geotechnische parameters. Stichting CUR, Gouda."),
        ("CUR-publicatie 166. Damwandconstructies. Stichting CUR, Gouda (4e/6e druk).",
         "CUR-publicatie 166 (2012). Damwandconstructies, deel 1 en 2, 6e druk. Stichting CUR, Gouda (incl. errata 2014)."),
        ("Blum, H. (1931). Einspannungsverhaeltnisse bei Bohlwerken.", "Blum, H. (1931). Einspannungsverhältnisse bei Bohlwerken."),
        ("Verruijt, A. (2012). Soil Mechanics. Delft University of Technology.",
         "Verruijt, A. (2012). Soil Mechanics. Delft University of Technology (2001, herzien 2012). https://geo.verruijt.net"),
        ("DIN 4150-3:2016 - Erschuetterungen im Bauwesen. Teil 3: Einwirkungen auf bauliche Anlagen.",
         "DIN 4150-3:2016-12 - Erschütterungen im Bauwesen. Teil 3: Einwirkungen auf bauliche Anlagen."),
        ("Bentley Systems (2024). PLAXIS 2D Reference Manual, Version 24. Bentley Systems, Delft.",
         "Seequent, The Bentley Subsurface Company (2024). PLAXIS 2D 2024 (V24, «gebruikte versie 2024.x in te vullen») Reference Manual. " + MANUALS_URL),
        ("Bentley Systems (2024). PLAXIS 2D Material Models Manual, Version 24. Bentley Systems, Delft.",
         "Seequent, The Bentley Subsurface Company (2024). PLAXIS 2D 2024 (V24, «gebruikte versie 2024.x in te vullen») Material Models Manual. " + MANUALS_URL),
        ("MADEP CPT (madep.be) - eigen interpretatietool sondeergegevens.",
         "MADEP CPT Interpreter (madep.be) - eigen interpretatietool sondeergegevens, «versie en datum in te vullen»."),
    ]
    for old, new in R:
        hit = False
        for p in ps[:70]:
            if replace_in_paragraph(p, old, new, "RN " + old[:40]):
                hit = True
                break
        if not hit:
            print("  ! Rekennota: not found:", old[:60])
    i = find(d, "Verificatie van de referenties:")
    replace_in_paragraph(ps[i], "dat laatste is CUR-publicatie 166.",
        "dat laatste is CUR-publicatie 166. De referentielijst werd op " + "29 augustus 2026" + " online geverifieerd (NBN-catalogus, Buildwise, uitgevers): "
        "de nationale bijlagen NBN EN 1990 ANB en NBN EN 1993-1-1 ANB, de sondeernorm NBN EN ISO 22476-1, de uitvoeringsnorm NBN EN 12063 en de "
        "BGGG-sondeerprocedure zijn naar de geldende edities gebracht; de titel van Infofiche 56.2 en de CUR 166-druk zijn gecorrigeerd.",
        "RN verification note")
    d.save(f)


def docx_to_md(src: Path, dst: Path):
    """Plain conversion for the worklog copies (paragraphs + pipe tables, in body order)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    d = docx.Document(src)
    out = []
    for el in d.element.body.iterchildren():
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            p = Paragraph(el, d)
            t = p.text.strip()
            if not t:
                continue
            s = p.style.name if p.style is not None else ''
            if s.startswith('Heading'):
                try:
                    lvl = int(s.split()[-1])
                except ValueError:
                    lvl = 1
                out.append('#' * lvl + ' ' + t)
            elif s == 'Title':
                out.append('# ' + t)
            else:
                out.append(t)
        elif tag == 'tbl':
            tb = Table(el, d)
            rows = [[c.text.replace('\n', ' ').strip() for c in r.cells] for r in tb.rows]
            if not rows:
                continue
            out.append('| ' + ' | '.join(rows[0]) + ' |')
            out.append('|' + '---|' * len(rows[0]))
            for r in rows[1:]:
                out.append('| ' + ' | '.join(r) + ' |')
        out.append('')
    dst.write_text('\n'.join(out), encoding='utf-8')


if __name__ == "__main__" and "--fixup" not in sys.argv:
    sheet_pile_manual()
    brinch_hansen_chapter()
    vibratory_chapter()
    rekennota()
    print(f"{len(changes)} replacements applied:")
    for c in changes:
        print("  -", c)
    text_dir = ROOT / "worklog" / "course-text"
    for name in ["Sheet_Pile_Retaining_Walls_Manual_EC7_PLAXIS_v24", "Brinch_Hansen_Tlat_Soldier_Pile_Walls_Textbook_Section",
                 "Vibratory_Pile_Installation_Manual_Eurocode_Course_Chapter", "Rekennota_beschoeiing_berlinerwand_HEA180"]:
        docx_to_md(COURSE / f"{name}.docx", text_dir / f"{name}.md")
    print("course-text regenerated")


# --------------------------------------------------------------------------------------------
# Fix-up (run after the first pass): the manual's "[working source]" tags are hyperlink elements
# (not runs), so the first pass duplicated them as plain text; and the hyperlinks themselves
# resolve to real Bentley KB articles (KB0110020, KB0109902, KB0109451) whose titles the manual
# had right — those entries are restored with their KB numbers instead of the Reference-Manual
# substitutes.
def sheet_pile_manual_fixup():
    f = COURSE / "Sheet_Pile_Retaining_Walls_Manual_EC7_PLAXIS_v24.docx"
    d = docx.Document(f)
    ps = d.paragraphs
    r0 = find(d, "References", start=200)
    ref = {int(ps[i].text.strip()[1:ps[i].text.index("]")]): ps[i]
           for i in range(r0 + 1, r0 + 18) if ps[i].text.strip().startswith("[")}
    for n, p in ref.items():
        for r in p.runs:
            if " [working source]" in r.text:
                r.text = r.text.replace(" [working source]", "")
                changes.append(f"SP[{n}] duplicate tag removed")
    replace_ref(ref[7], 7, "Seequent / Bentley Systems. PLAXIS Knowledge Base, plate data sets for sheet piles: “Material datasets for plates: "
                "sheet pile wall in bending” (KB0110039, formerly wiki 45929; last modified 11 July 2025) and “Material parameter datasets for "
                "sheetpiles and beams” (formerly wiki 46023); worked sheet-pile Plate data set (linked article KB0051314). ", "SP[7] KB titles")
    replace_ref(ref[8], 8, "Seequent / Bentley Systems. Modelling soil-structure interaction: interfaces. PLAXIS Knowledge Base article KB0110020 "
                "(PLAXIS 2D/3D, created 20 August 2012). ", "SP[8] restored + KB")
    replace_ref(ref[10], 10, "Seequent / Bentley Systems. Permeability in interfaces: Practical situations. PLAXIS Knowledge Base article KB0109902 "
                "(Tips and tricks, PLAXIS 2D/3D, created 4 July 2017). ", "SP[10] restored + KB")
    replace_ref(ref[11], 11, "Seequent / Bentley Systems. Permeability in interfaces. PLAXIS Knowledge Base article KB0109451: definition of hydraulic "
                "resistance d/k and drainage conductivity d·k (also in the PLAXIS 2D Reference Manual, “Interfaces tabsheet”). ", "SP[11] restored + KB")
    d.save(f)


if __name__ == "__main__" and "--fixup" in sys.argv:
    changes.clear()
    sheet_pile_manual_fixup()
    print(f"fix-up: {len(changes)} replacements:")
    for c in changes:
        print("  -", c)
    name = "Sheet_Pile_Retaining_Walls_Manual_EC7_PLAXIS_v24"
    docx_to_md(COURSE / f"{name}.docx", ROOT / "worklog" / "course-text" / f"{name}.md")
